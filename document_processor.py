import os
import fitz  # PyMuPDF
import docx
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter

def extract_table_from_pdf(page):
    """Extract tables from a PDF page and convert to a text representation."""
    tables = page.find_tables()
    table_text = ""
    
    for table in tables:
        try:
            # Convert to pandas DataFrame
            df = table.to_pandas()
            # Format as string with alignment
            table_text += "\n" + df.to_string(index=False) + "\n"
        except Exception as e:
            table_text += f"\nTable extraction error: {str(e)}\n"
    
    return table_text

def process_pdf(file_path):
    """Process a PDF file and extract text with tables, preserving structure."""
    try:
        doc = fitz.open(file_path)
        file_name = os.path.basename(file_path)
        all_text = []
        
        for page_num, page in enumerate(doc):
            # Extract regular text
            text = page.get_text()
            
            # Extract tables
            table_text = extract_table_from_pdf(page)
            
            # Combine text and tables
            if table_text:
                combined_text = text + "\n\nTABLE CONTENT:\n" + table_text
            else:
                combined_text = text
            
            # Add document metadata
            metadata = {
                "source": file_name,
                "page": page_num + 1
            }
            
            all_text.append({"text": combined_text, "metadata": metadata})
        
        return all_text
    
    except Exception as e:
        print(f"Error processing PDF {file_path}: {str(e)}")
        return []

def process_docx(file_path):
    """Process a DOCX file and extract text with tables, preserving structure."""
    try:
        doc = docx.Document(file_path)
        file_name = os.path.basename(file_path)
        
        # Extract paragraphs
        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Extract tables
        tables_content = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(cells)
            
            if rows:
                # Convert to DataFrame for proper formatting
                df = pd.DataFrame(rows)
                
                # If the first row looks like a header
                if len(df) > 1:
                    headers = df.iloc[0].tolist()
                    df.columns = headers
                    df = df.iloc[1:]
                
                tables_content.append(f"\nTABLE {i+1}:\n{df.to_string(index=False)}")
        
        # Combine all content
        full_content = "\n\n".join(text_content + tables_content)
        
        # Add metadata
        result = [{
            "text": full_content,
            "metadata": {
                "source": file_name,
                "page": 1  # For simplicity in DOCX files
            }
        }]
        
        return result
    
    except Exception as e:
        print(f"Error processing DOCX {file_path}: {str(e)}")
        return []

def test_document_processing():
    """Test function to verify document processing works correctly."""
    # Make sure the docs directory exists
    os.makedirs("docs", exist_ok=True)
    
    # Check if there are any documents
    pdf_files = [f for f in os.listdir("docs") if f.lower().endswith(".pdf")]
    docx_files = [f for f in os.listdir("docs") if f.lower().endswith(".docx")]
    
    if not pdf_files and not docx_files:
        print("No documents found in the 'docs' directory.")
        return
    
    # Process one PDF and one DOCX for testing
    if pdf_files:
        pdf_path = os.path.join("docs", pdf_files[0])
        print(f"Testing PDF processing with {pdf_path}")
        pdf_results = process_pdf(pdf_path)
        
        if pdf_results:
            print(f"Successfully processed PDF: {len(pdf_results)} pages")
            # Print sample of first page
            first_page = pdf_results[0]["text"]
            print(f"Sample of first page (first 200 chars):\n{first_page[:200]}...")
        else:
            print("Failed to process PDF")
    
    if docx_files:
        docx_path = os.path.join("docs", docx_files[0])
        print(f"Testing DOCX processing with {docx_path}")
        docx_results = process_docx(docx_path)
        
        if docx_results:
            print(f"Successfully processed DOCX: {len(docx_results)} sections")
            # Print sample
            first_section = docx_results[0]["text"]
            print(f"Sample of content (first 200 chars):\n{first_section[:200]}...")
        else:
            print("Failed to process DOCX")

def test_text_splitting():
    """Test the text splitter functionality."""
    sample_text = """
    This is a test document with multiple paragraphs.
    
    It has several sections that should be split appropriately.
    
    TABLE CONTENT:
       Column A    Column B    Column C
    1    Value 1    Value 2    Value 3
    2    Value 4    Value 5    Value 6
    
    The text continues after the table with more information.
    Some of this information might be related to the table.
    
    Another paragraph follows with different content.
    """
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=100,
        chunk_overlap=20,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = text_splitter.split_text(sample_text)
    
    print(f"Split text into {len(chunks)} chunks:")
    for i, chunk in enumerate(chunks):
        print(f"--- Chunk {i+1} ---")
        print(chunk)
        print("-----------------")

if __name__ == "__main__":
    test_document_processing()
    test_text_splitting()