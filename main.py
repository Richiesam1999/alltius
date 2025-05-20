import os
import re
import glob
import gradio as gr
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing
import pandas as pd
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
from langchain.llms import HuggingFaceHub
import torch

# Check if we're running in a GPU environment
device = "cuda" if torch.cuda.is_available() else "cpu"
print(f"Using device: {device}")

# Create directories if they don't exist
os.makedirs("docs", exist_ok=True)
os.makedirs("db", exist_ok=True)

# Function to extract text from PDF files with page numbers
def extract_from_pdf(pdf_path):
    documents = []
    pdf_document = fitz.open(pdf_path)
    file_name = os.path.basename(pdf_path)
    
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        text = page.get_text()
        
        # Handle tables by extracting them as text
        tables = page.find_tables()
        table_text = ""
        for table in tables:
            df = table.to_pandas()
            table_text += df.to_string() + "\n\n"
        
        # Combine regular text and table text
        combined_text = text
        if table_text:
            combined_text += "\nTABLE CONTENT:\n" + table_text
            
        if combined_text.strip():  # Only add non-empty pages
            metadata = {
                "source": file_name,
                "page": page_num + 1  # 1-indexed page number for human readability
            }
            documents.append(Document(page_content=combined_text, metadata=metadata))
    
    return documents

# Function to extract text from DOCX files
def extract_from_docx(docx_path):
    documents = []
    doc = docx.Document(docx_path)
    file_name = os.path.basename(docx_path)
    
    full_text = []
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    
    # Extract tables
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        
        # Convert table to DataFrame then to string
        if table_data:
            df = pd.DataFrame(table_data)
            # Use the first row as header if it looks like a header
            if len(df) > 1:
                df.columns = df.iloc[0]
                df = df[1:]
            table_text = f"\nTABLE {i+1}:\n" + df.to_string(index=False)
            full_text.append(table_text)
    
    # Combine all text with metadata
    combined_text = "\n".join(full_text)
    if combined_text.strip():
        metadata = {
            "source": file_name,
            # Note: For DOCX files, we can't easily get page numbers
            # but we'll set it to 1 for consistency in the database
            "page": 1
        }
        documents.append(Document(page_content=combined_text, metadata=metadata))
    
    return documents

# Function to process documents and create chunks
def process_documents():
    documents = []
    
    # Process all PDF files
    for pdf_file in glob.glob("docs/*.pdf"):
        documents.extend(extract_from_pdf(pdf_file))
    
    # Process all DOCX files
    for docx_file in glob.glob("docs/*.docx"):
        documents.extend(extract_from_docx(docx_file))
    
    # Create chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    chunks = text_splitter.split_documents(documents)
    print(f"Created {len(chunks)} chunks from {len(documents)} document sections")
    
    return chunks

# Function to create embeddings and store in Chroma
def create_embeddings(chunks):
    # Use Nomic embeddings (Nomic-Embed-Text-v1)
    embeddings = HuggingFaceEmbeddings(
    model_name="nomic-ai/nomic-embed-text-v1",
    model_kwargs={
        'device': device,
        'trust_remote_code': True  
    }
)
    # Create and persist the vector store
    db = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory="db"
    )
    db.persist()
    
    return db

# Function to load existing embeddings
def load_embeddings():
    embeddings = HuggingFaceEmbeddings(
        model_name="nomic-ai/nomic-embed-text-v1",
        model_kwargs={
        'device': device,
        'trust_remote_code': True}
    )
    
    # Load the persisted vector store
    db = Chroma(
        persist_directory="db",
        embedding_function=embeddings
    )
    
    return db

# Query function to get relevant information
def query_documents(query_text, db, k=3):
    # Get relevant documents
    docs = db.similarity_search_with_score(query_text, k=k)
    
    # Format results with sources
    results = []
    for doc, score in docs:
        source = doc.metadata.get("source", "Unknown")
        page = doc.metadata.get("page", "Unknown")
        results.append({
            "content": doc.page_content,
            "source": f"{source} (Page {page})",
            "relevance_score": f"{score:.2f}"
        })
    
    return results

# Function to format results for display
def format_results(results):
    if not results:
        return "No relevant information found."
    
    formatted_text = ""
    for i, result in enumerate(results, 1):
        formatted_text += f"Result {i} [Source: {result['source']}]\n"
        formatted_text += f"Relevance: {result['relevance_score']}\n"
        formatted_text += f"{result['content']}\n\n"
        formatted_text += "-" * 80 + "\n\n"
    
    return formatted_text

# Main function for query answering
def answer_query(query, db, use_llm=False):
    if use_llm:
        # If using LLM-based answers (optional)
        # This section would integrate with an LLM for generating answers
        # For now, we'll just retrieve and format the relevant chunks
        pass
    
    # Get relevant documents
    results = query_documents(query, db)
    
    # Format results
    response = format_results(results)
    
    return response

# Gradio Interface
def create_gradio_interface(db):
    with gr.Blocks(title="Document RAG System") as demo:
        gr.Markdown("# Document RAG System")
        gr.Markdown("Ask questions about your documents and get relevant information with sources.")
        
        with gr.Row():
            with gr.Column(scale=3):
                query_input = gr.Textbox(label="Your Question", placeholder="What is the overall deductible?")
                submit_button = gr.Button("Submit")
                
            with gr.Column(scale=1):
                num_results = gr.Slider(minimum=1, maximum=10, value=3, step=1, label="Number of Results")
                # Placeholder for future file upload feature
                # file_output = gr.File(label="Upload Document")
        
        response_output = gr.Textbox(label="Response with Sources", lines=15)
        
        submit_button.click(
            fn=lambda query, k: answer_query(query, db, use_llm=False),
            inputs=[query_input, num_results],
            outputs=response_output
        )
    
    return demo

# Main function to run the application
def main():
    # Check if database already exists
    if os.path.exists("db") and os.listdir("db"):
        print("Loading existing database...")
        db = load_embeddings()
    else:
        print("Creating new database...")
        # Process documents and create embeddings
        chunks = process_documents()
        db = create_embeddings(chunks)
    
    # Create and launch Gradio interface
    demo = create_gradio_interface(db)
    demo.launch()

if __name__ == "__main__":
    main()